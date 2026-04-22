from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read markdown
with open('C:/Users/kajiy/Downloads/trustsystem_recruit/WORK_REPORT.md', 'r', encoding='utf-8') as f:
    md_content = f.read()

# Create docx
doc = Document()
lines = md_content.split('\n')
in_table = False
current_table = None

for line in lines:
    line = line.rstrip()

    # Skip empty lines
    if not line:
        if in_table:
            in_table = False
        doc.add_paragraph()
        continue

    # Headers
    if line.startswith('# '):
        p = doc.add_heading(line[2:], level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)

    # Metadata
    elif ':' in line and '**' not in line and not line.startswith('|'):
        if '：' in line:
            parts = line.split('：', 1)
        else:
            parts = line.split(':', 1)
        if len(parts) == 2:
            p = doc.add_paragraph()
            run = p.add_run(parts[0].strip() + '：')
            run.bold = True
            p.add_run(' ' + parts[1].strip())

    # Tables
    elif line.startswith('|'):
        if not in_table:
            in_table = True
            cols = [cell.strip() for cell in line.split('|')[1:-1]]
            current_table = doc.add_table(rows=1, cols=len(cols))
            current_table.style = 'Light Grid Accent 1'
            header_cells = current_table.rows[0].cells
            for i, col in enumerate(cols):
                header_cells[i].text = col
        elif '---|' not in line:
            cols = [cell.strip() for cell in line.split('|')[1:-1]]
            row_cells = current_table.add_row().cells
            for i, col in enumerate(cols):
                row_cells[i].text = col

    # Code blocks
    elif line.startswith('```'):
        continue

    # Bullet points
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')

    # Normal text
    elif line and not line.startswith(' '):
        doc.add_paragraph(line)

# Save
doc.save('C:/Users/kajiy/Downloads/trustsystem_recruit/WORK_REPORT.docx')
print('WORK_REPORT.docx created successfully!')
